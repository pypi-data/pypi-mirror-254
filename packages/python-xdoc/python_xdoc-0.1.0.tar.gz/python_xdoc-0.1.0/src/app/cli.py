import os
import re
import yaml
import pytz
import uuid
import httpx
import shutil
import string
import random
import qrcode
import logging
import hashlib
import getpass
import asyncio
import requests
import argparse
import traceback
import mimetypes
import tracemalloc
import gradio as gr
import uvloop

from mega import Mega
from tqdm import tqdm
from io import BytesIO
from pyppeteer import launch
from bs4 import BeautifulSoup
from rich.console import Console
from markdown_it import MarkdownIt
from Cryptodome.Hash import SHA256
from PIL import Image, ImageFont, ImageDraw
from pypdf import PdfMerger, PdfWriter, PdfReader

from modules.api import WebAPI
from modules.common import Interface

from docx import Document
from docx.shared import Pt, Mm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml

asyncio.set_event_loop_policy(uvloop.EventLoopPolicy())

class PdfProcessor:

    def __init__(self, params):
        self.params = params

    async def gen(self):
        output_file = os.path.abspath(self.params['output'][0])

        console.print(f'[*] self.params: {self.params}')
        html_dir = os.path.dirname(self.params['input'][0])
        os.chdir(html_dir)
        browser = await launch(
            args=[
            '--headless',
            ]
        )
        page = await browser.newPage()
        merger = PdfMerger()

        md = MarkdownIt()
        temp_html_file = self.params['tmp_file'] + '.html'
        tmp_dir = self.params['tmp_dir']
        try:
            for index, input_file in enumerate(self.params['input']):
                if input_file.lower().endswith(".md"):
                    with open(input_file, "r", encoding="utf-8") as file:
                        markdown_data = file.read()
                        html_data = md.render(markdown_data)
                else:
                    with open(input_file, "r", encoding="utf-8") as file:
                        html_data = file.read()

                await page.setContent(html_data)

                css_file = self.params.get('css')
                if css_file:
                    if css_file.startswith('http://') or css_file.startswith('https://'):
                        css_link = f'<link rel="stylesheet" type="text/css" href="{css_file}">\n'
                        html_data = css_link + html_data
                        await page.goto(css_file)
                    elif os.path.isfile(css_file):
                        with open(css_file, "r", encoding="utf-8") as css_file:
                            css_content = css_file.read()
                            css_style = f'<style>\n{css_content}\n</style>\n'
                            html_data = css_style + html_data
                        await page.addStyleTag({'path': os.path.abspath(css_file)})
                    else:
                        # css_content = css_file.read() # TODO
                        css_style = f'<style>\n{css_file}\n</style>\n'
                        html_data = css_style + html_data
                        await page.addStyleTag({'content': css_file})

                await page.emulateMedia(self.params['emulateMedia'])

                with open(temp_html_file, "w", encoding="utf-8") as file:
                    file.write(html_data)

                pdf_data = await page.pdf(self.params)

                pdf_buffer = BytesIO(pdf_data)

                tasks = await ImageProcessor.img_link(input_file, markdown_data, tmp_dir)
                for image_path in tasks:
                    pass
                    # await ImageProcessor.add_img(pdf_buffer, image_path) # TODO
                merger.append(pdf_buffer)

                # if os.path.exists(temp_html_file):
                #     os.remove(temp_html_file)

            if self.params.get('meta'):
                merger.add_meta(self.params.get('meta'))

            return output_file

        except Exception as e:
            console.print(f'[!] {e}')

        finally:
            await browser.close()
            merger.write(output_file)

            merger.close()
            console.print(f'[i] Temporally html data : {temp_html_file}')
            console.print(f'[*] Output : {output_file}')

    def merge(self):
        try:
            merger = PdfMerger()
            for file in input_files:
                merger.append(file)
            with open(output_file, 'wb') as merged:
                merger.write(merged)
            console.print('[*] Merge operation completed.')
        except Exception as e:
            console.print(f'[!] {e}')

    def write(self, start_page, end_page):
        input_file = self.params['input'][0]
        output_file = self.params['output'][0]
        try:
            reader = PdfReader([input_file])
            writer = PdfWriter()

            for page_num in range(start_page - 1, end_page):
                writer.add_page(reader.pages[page_num])

            with open(output_file, 'wb') as extracted:
                writer.write(extracted)
            console.print('[*] Extract operation completed.')

        except Exception as e:
            console.print(f'[!] {e}')

    def split(self, output_prefix, split_page):
        input_file = self.params['input'][0]
        reader = PdfReader(input_file)
        try:
            for page_num in range(len(reader.pages)):
                writer = PdfWriter()
                writer.add_page(reader.pages[page_num])

                output_file = f"{output_prefix}_{page_num + 1}.pdf"
                with open(output_file, 'wb') as output:
                    writer.write(output)

            console.print('[*] Split operation completed.')
        except Exception as e:
            console.print(f'[!] {e}')

    def extract_text(self, output_text_file):
        input_file = self.params['input'][0]
        reader = PdfReader(input_file)
        text = ""
        try:
            for page_num in range(len(reader.pages)):
                text += reader.pages[page_num].extractText()
            with open(output_text_file, 'w', encoding='utf-8') as text_file:
                text_file.write(text)

            console.print('[*] Text extraction operation completed.')
        except Exception as e:
            console.print(f'[!] {e}')

    def select_meta(self, selected_meta):
        input_file = self.params['input'][0]
        output_file = self.params['output'][0]
        try:
            reader = PdfReader(input_file)
            writer = PdfWriter()

            for page_num in range(len(reader.pages)):
                writer.add_page(reader.pages[page_num])

            writer.removeMetadata(selected_meta)

            with open(output_file, 'wb') as output:
                writer.write(output)

            console.print('[*] Selective meta removal operation completed.')
        except Exception as e:
            console.print(f'[!] {e}')

    def remove_meta(self):
        input_file = self.params['input'][0]
        try:
            reader = PdfReader(input_file)
            writer = PdfWriter()

            for page_num in range(len(reader.pages)):
                writer.add_page(reader.pages[page_num])

            writer.removeMetadata()

            with open(output_file, 'wb') as output:
                writer.write(output)

            console.print('[*] Remove meta operation completed.')
        except Exception as e:
            console.print(f'[!] {e}')

    def encrypt(self, not_yet, password):
        # print work dir
        encrypted_file = self.params['output'][0]

        console.print(f'[*] : {encrypted_file}')
        reader = PdfReader(open((encrypted_file), 'rb'))
        writer = PdfWriter()
        writer.encrypt(password)

        for page_number in range(len(reader.pages)):
            page = reader.pages[page_number]
            writer.add_page(page)
        with open(encrypted_file, 'wb') as output:
            writer.write(output)

    def decrypt(self, pwd):
        reader = PdfReader(self.params['input'][0])
        writer = PdfWriter()
        if reader.is_encrypted:
            reader.decrypt(pwd)
        for page in reader.pages:
            writer.add_page(page)

        with open(self.params['output'][0], "wb") as f:
            writer.write(f)

class Encryptor:

    def gen_pass(length=8):
        characters = string.ascii_letters + string.digits + string.punctuation
        password = ''.join(random.choice(characters) for _ in range(length))
        return password

    def hash_pass(password, salt):
        key = SHA256.new(password.encode('utf-8') + salt).digest()
        return key

class ImageProcessor:

    async def gen_qr(data, out_path, **qr_params):
        try:
            qr = qrcode.QRCode(**qr_params)
            qr.add_data(data)
            img = qr.make_image(fill_color="black", back_color="white")
            box_size = int(qr_params['box_size'])
            img = img.resize((box_size, box_size))
            img.save(out_path, 'PNG')
            return qr
        except Exception as e:
            console.print(f"[!] Failed to generate QR code for {data}: {str(e)}")
            traceback.print_exc()
            return None

    async def add_qr(input_files, qr_params, tmp_dir, tmp_file):
        try:
            for input_file in input_files:
                temp_md_file = f'{tmp_file}.md'
                if input_file.lower().endswith(".md"):
                    with open(input_file, "r") as file:
                        md_doc = file.read()

                        pattern = re.compile(r'\[(.*?)\]\((https?://[^\s\)]+|[^)]+)\)')
                        links = pattern.findall(md_doc)
                        tasks = []
                        for idx, (link_text, url) in enumerate(links):
                            qr_file = f"qr_{idx}.png" # TODO
                            qr_path = os.path.join(tmp_dir, f"qr_{idx}.png")
                            task = asyncio.ensure_future(ImageProcessor.gen_qr(link_text, qr_path, **qr_params))
                            tasks.append((link_text, url, qr_file, task))

                        for link_text, url, qr_file, task in tasks:
                            qr = await task
                            if qr:
                                qr_link = f"![]({qr_file})"
                                md_doc = md_doc.replace(f"[{link_text}]({url})", f"<div class='qr'>\n\n[{link_text}]({url}) {qr_link}\n</div>")

                    with open(temp_md_file, "w") as file:
                        file.write(md_doc)

            return tmp_dir
            console.print(f'[i] Temporary markdown docs in: {temp_md_file}')
            console.print(f'[i] QRs in: {tmp_dir}')
        except Exception as e:
            console.print(f'[!] {e}')

    async def add_img(pdf_page, image_path):
        c = canvas.Canvas(pdf_page)
        img = ImageReader(image_path)
        c.drawImage(img, 100, 100)
        c.showPage()
        c.save()

    @staticmethod
    async def img_link(einput_file, md_doc, tmp_dir):
        pattern = re.compile(r'\[(.*?)\]\((https?://[^\s\)]+|[^)]+)\)')
        links = pattern.findall(md_doc)

        image_mime_types = ["image/jpeg", "image/png", "image/gif", "image/bmp", "image/webp"]

        tasks = []
        try:
            async with httpx.AsyncClient() as client:
                for idx, (link_text, url) in enumerate(links):
                    image_file = f"img_{idx}.png"
                    image_path = os.path.abspath(os.path.join(tmp_dir, image_file))
                    if url.startswith('http://') or url.startswith('https://'):
                        response = await client.head(url)
                        if response.status_code == 200:
                            content_type = response.headers.get("content-type", "").lower()
                            if any(content_type.startswith(mime) for mime in image_mime_types):
                                response = await client.get(url)
                                image_data = response.content
                                with open(image_path, "wb") as file:
                                    file.write(image_data)
                                tasks.append(image_path)
                            else:
                                console.print(f"[!] URL {idx} is not an image.")
                        else:
                            console.print(f"[!] Cannot fetch image file {idx}.")
                    elif os.path.isfile(url):
                        # url = os.path.dirname(url)
                        # image_path = os.path.dirname(image_path)
                        console.print(f'[i] Copy {url} to {image_path}')
                        shutil.copy(url, image_path)
                        tasks.append(url)
                # console.print(f'tasks: {tasks}')
            return tasks

        except Exception as e:
            console.print(f'[!] {e}')

class DocxProcessor:

    def __init__(self, params):
        self.params = params

    def add_hyperlink(paragraph, url, text, color, underline):
        """
        A function that places a hyperlink within a paragraph object.

        :param paragraph: The paragraph we are adding the hyperlink to.
        :param url: A string containing the required url
        :param text: The text displayed for the url
        :return: The hyperlink object
        """

        part = paragraph.part
        r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, nsdecls('r'))
        new_run = docx.oxml.shared.OxmlElement('w:r')
        rPr = docx.oxml.shared.OxmlElement('w:rPr')
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)
        r = paragraph.add_run ()
        r._r.append (hyperlink)
        r.font.color.rgb = color
        r.font.underline = underline

        return hyperlink

    def calculate_max_size(img, max_width, max_height):
        aspect_ratio = img.width / img.height
        new_height = min(max_height, img.height)
        new_width = min(new_height * aspect_ratio, img.width)

        return new_width, new_height

    def resize_image(img, max_width, max_height):
        new_width, new_height = int(max_width), int(max_height)
        if new_width > 0 and new_height > 0:
            img = img.resize((new_width, new_height))

        return img

    def set_page_title(image_path):
        file_name = os.path.basename(image_path)
        page_title = file_name.split('-')[1].strip()
        return page_title

    def create_element(name):
        return OxmlElement(name)

    def create_attribute(element, name, value):
        element.set(ns.qn(name), value)

    def add_page_number(run):
        fldChar1 = create_element('w:fldChar')
        create_attribute(fldChar1, 'w:fldCharType', 'begin')

        instrText = create_element('w:instrText')
        create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE / NUMPAGES"

        fldChar2 = create_element('w:fldChar')
        create_attribute(fldChar2, 'w:fldCharType', 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    def extract_information_from_readme(readme_path):
        with open(readme_path, 'r', encoding='utf-8') as readme_file:
            readme_content = readme_file.read()

        md_content = markdown(readme_content)

        title_start = md_content.find('<!-- title:')
        title_end = md_content.find('-->', title_start)

        if title_start != -1 and title_end != -1:
            page_title = md_content[title_start + len('<!-- title:'):title_end].strip()
            return page_title
        else:
            return ""

    def create_docx(input_dir, template_file, output_file, max_width, max_height, font_path, readme_enabled):
        readme_path = os.path.join(input_dir, 'readme.md') if readme_enabled else None
        ts, tmpd = Interface.get_tempdir()
        template_doc = Document(template_file)
        doc = Document()
        image_files = [f for f in os.listdir(input_dir) if f.endswith('.png')]
        total_images = len(image_files)

        # Meta
        doc.core_properties.author = "Admin"
        doc.core_properties.title = "My Document"
        doc.core_properties.subject = "The subject of my document"
        doc.core_properties.created = ts

        total_lines = 26
        for i in range(int(total_lines/2) - 1):
            doc.add_paragraph()

        # Title
        title = doc.add_paragraph()
        title_run = title.add_run('Title')
        title_run.font.name = font_path
        title_run.bold = True
        title_run.font.size = Pt(40)
        title.paragraph_format.alignment = 1

        doc.add_page_break()

        # Patagraph
        for paragraph in doc.paragraphs:
            paragraph.paragraph_format.keep_with_next = True
            paragraph.paragraph_format.widow_control = True
            paragraph.paragraph_format.alignment = 1
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.left_indent = Pt(0)
            paragraph.paragraph_format.right_indent = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing_rule = 0
            paragraph.paragraph_format.width = Pt(max_width / 2)

        # TOC
        toc = doc.add_paragraph()
        toc_index = toc.add_run('TOC\n')
        toc_index.bold = False
        toc_index.font.size = Pt(20)
        toc.style = 'TOC Heading'

        for index, image_file in enumerate(tqdm(image_files, desc="Processing Pages", unit="pp", total=total_images)):
            run = toc.add_run(f'{index + 1}. {image_file}')
            run.font.size = Pt(8)
            run.add_break()

        for index, image_file in enumerate(tqdm(image_files, desc="Processing Pages", unit="pp", total=total_images)):
            # Header
            paragraph = doc.sections[0].header.paragraphs[0]
            if len(paragraph.runs) > 0:
                paragraph.runs[0].text = f'{image_file}'
                paragraph.runs[0].font.size = Pt(20)
            else:
                subtitle = paragraph.add_run(f'{image_file}')
                subtitle.font.size = Pt(20)

            if readme_enabled:
                page_title = extract_information_from_readme(readme_path)
            else:
                page_title = ""

            # Image
            img_path = os.path.join(input_dir, image_file)
            img = Image.open(img_path)
            max_width, max_height = calculate_max_size(img, max_width, max_height)
            img = resize_image(img, max_width, max_height)

            if img is None:
                continue

            if index > 0:
                doc.add_page_break()

            img_stream = BytesIO()
            img.save(img_stream, format='PNG')
            doc.add_picture(img_stream)
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            subtitle.text = f'({page_title})'

            caption = f"Fig. {index + 1}: {image_file}"
            caption_paragraph = doc.add_paragraph(caption)
            caption_run = caption_paragraph.runs[0]
            caption_run.font.size = Pt(10)
            caption_paragraph.paragraph_format.alignment = 1

            # Body
            doc.add_section()

        # Footer
        doc.add_section(WD_SECTION.NEW_PAGE)
        doc.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # doc.sections[-1].footer.is_linked_to_previous = False
        add_page_number(doc.sections[-1].footer.paragraphs[0].add_run())

        # Description
        page_description = doc.add_paragraph('Page numbers are indicated in the TOC.')
        doc.add_page_break()

        for element in template_doc.element.body:
            doc.element.body.append(element)

        # Output
        output_dir = os.path.dirname(output_file)

        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        doc.save(output_file)
        print(f'Done at {output_dir}')

class Parser:

    def input_txt(files):
        ts, tmp_dir = Interface.get_tempdir()
        input_files = []

        for index, file in enumerate(files):
            mime_type, _ = mimetypes.guess_type(file)
            if mime_type and mime_type.startswith('text'):
                try:
                    if file.startswith("http://") or file.startswith("https://"):
                        response = requests.get(file)
                        if response.status_code == 200:
                            soup = BeautifulSoup(response.text, 'html.parser')
                            content = soup.get_text()
                            filename = os.path.join(tmp_dir, os.path.basename(file))
                            with open(filename, 'w', encoding='utf-8') as md_file:
                                md_file.write(content)

                            input_files.append(filename)
                            console.print(f"[i] Temporary file from '{file}' to {filename}")
                        else:
                            console.print(f"[!] to retrieve content from {file}")
                    else:
                        input_files.append(os.path.abspath(file))
                except Exception as e:
                    console.print(f"[!] {e}")
            elif mime_type and mime_type.startswith('application'):
                console.print(f"[i] File {index} is application/* file.")
                input_files.append(os.path.abspath(file))
            else:
                console.print(f"[!] File {index} is not a human-readable.")
                input_files.append(os.path.abspath(file))

        return input_files

    def margin(value):
        parts = value.split()
        result = {}
        i = 0
        key = None

        for i in range(len(parts)):
            if not re.match(r'\d+', parts[i]):
                key = parts[i]

            for j in range(i + 1, len(parts)):
                if re.match(r'\d+', parts[j]):
                    if key is not None and key not in result:
                        result[key] = parts[j]
                    break

        return result

    def meta(value):
        parts = value.split()
        result = {}
        current_key = None

        for part in parts:
            if current_key is None:
                current_key = part
            else:
                result[current_key] = part
                current_key = None

        return result

    def template(template_str):
        if template_str.startswith("http://") or template_str.startswith("https://"):
            response = requests.get(template_str)
            if response.status_code == 200:
                return response.text
            else:
                raise ValueError("[!] Failed to fetch template from URL.")

        elif os.path.isfile(template_str):
            with open(template_str, "r", encoding="utf-8") as file:
                return file.read()

        else:
            return template_str

def choose_file(input_files):
    while True:
        console.print("Which file would you like to process?")
        for i, input_file in enumerate(input_files):
            console.print(f"{i + 1}. {input_file}")

        choice = input()
        try:
            choice = int(choice)
            if 1 <= choice <= len(input_files):
                return input_files[choice - 1]
            else:
                console.print("[!] Invalid choice. Please select a valid file.")
        except ValueError:
            console.print("[!] Invalid input. Please enter a valid number.")

def clear_directory():
    try:
        dir = "/tmp/"
        for root, dirs, files in os.walk(dir, topdown=False):
            for name in files:
                os.remove(os.path.join(root, name))
        gr.Info(f'Clear cache done.')
    except Exception as e:
        gr.Error(f'Failed to delete files in {dir}. Reason: {e}')

async def webui(params, usage_doc):
    async def wrapper(*params):
        files, styles, header, footer, exec_clear = params
        if exec_clear:
            clear_directory()

        ts, tmp_dir = Interface.get_tempdir()
        tmp_file = f'{tmp_dir}/{ts}'
        os.makedirs(tmp_dir, exist_ok=True)

        input_files = [os.path.abspath(file.name) for file in files]
        output_file = tmp_file + '.pdf'

        params = {
            'emulateMedia': 'screen',
            'scale': 1,
            'format': 'a4',
            'css': styles.name if styles else None,
            'headerTemplate': header.name if header else None,
            'footerTemplate': footer.name if footer else None,
        }

        params.update({'input': input_files})
        params.update({'output': output_file})
        params.update({'tmp_dir': tmp_dir, 'tmp_file': f'{tmp_dir}/{ts}'})

        pdf_proc = PdfProcessor(params)

        processing_file = await pdf_proc.gen()

        gr.Info(f'[*] Output to: {output_file}')

        return gr.Files(processing_file)


    version = 'v0.1.0'
    css = ''
    with open('src/app/css/style.css', 'r') as f:
        css = f.read()
    with gr.Blocks(
        css=css, title='DocMaker Web UI'
        ) as app:
        gr.HTML(f"<h1>DocMaker Web UI {version}</h1>")
        with gr.Row(variant='panel'):
            clr_cache_btn = gr.ClearButton(value='C')
            run_btn = gr.Button(value='▶', variant="primary")

        with gr.Tab('Processor') as tab1:
            with gr.Row():
                with gr.Column(variant='panel') as col1:
                    input_file = gr.Files(label="demo.md", elem_id=f'{uuid.uuid4()}')
                    with gr.Accordion("Style", open=False):
                        input_style = gr.File(label="style.css", elem_id=f'{uuid.uuid4()}')
                    with gr.Accordion("Template", open=False):
                        input_header_html = gr.File(label="header.html", elem_id=f'{uuid.uuid4()}')
                        input_footer_html = gr.File(label="footer.html", elem_id=f'{uuid.uuid4()}')

                with gr.Column(variant='panel'):
                    output_file = gr.Files(elem_id=f'{uuid.uuid4()}')

            inputs = [
                input_file,
                input_style,
                input_header_html,
                input_footer_html,
                clr_cache_btn
            ]

            examples = [
                [['src/md/invitation/demo_1.md'], 'src/css/style_02.css', 'src/html/template/header.html', 'src/html/template/footer.html'],
                [None, None, None, None],
            ]
            gr.Examples(examples, inputs)

        with gr.Tab('Help') as tab2:
            with gr.Row():
                gr.Markdown(f'''
                # Convert any document format to PDF format.

                ### CLI Usage

                ```bash
                {usage_doc}
            ```
                ''')

        outputs = [
            output_file,
        ]

        run_btn.click(wrapper, inputs, outputs)

    app.queue().launch(
        # inline=True,
        inbrowser=False,
        # share=True,
        # debug=True,
        # max_threads=40,
        # auth=('test','pass'),
        # auth_message=auth_message,
        # prevent_thread_lock=False,
        show_error=True,
        server_name=None,
        # server_port=None,
        # show_tips=False,
        # height=,
        # width=,
        # encrypt=,
        # favicon_path="assets/favicon.ico",
        ssl_keyfile=None,
        ssl_certfile=None,
        ssl_keyfile_password=None,
        ssl_verify=True,
        quiet=False,
        # show_api=False,
        # file_directories=False,
        # allowed_paths=False,
        # blocked_paths=False,
        # root_path=False,
        # app_kwargs=False
        )

console = Console()

async def main():
    tracemalloc.start()
    logger = logging.getLogger("httpx")
    logger.setLevel(logging.ERROR)

    parser = argparse.ArgumentParser(description="Generate PDF from HTML pages.")

    parser.add_argument("-i", "--input", nargs="+", help="Input files.")
    parser.add_argument("-o", "--output", nargs="+", help="Output file path.")
    parser.add_argument("--scale", type=float, help="Scale of the printed page | 0.8 for 80 percent.")
    parser.add_argument("--displayHeaderFooter", action="store_true", help="Display header and footer.")
    parser.add_argument( "--headerTemplate", type=Parser.template, help="Add header template as [url|path|content].", )
    parser.add_argument( "--footerTemplate", type=Parser.template, help="Add footer template as [url|path|content].", )
    parser.add_argument("--printBackground", action="store_true", help="Print background images.")
    parser.add_argument("--landscape", action="store_true", help="Use landscape paper orientation.")
    parser.add_argument("--pageRanges", type=str, help="Page ranges to print | '1-5,8,11-13'.")
    parser.add_argument("--format", type=str, help="Paper format | 'A4'.")
    parser.add_argument("--width", type=str, help="Paper width | '10cm'.")
    parser.add_argument("--height", type=str, help="Paper height | '10mm'.")
    parser.add_argument("--margin", type=Parser.margin, help="Margin values (e.g., 'top 10mm right 10mm')")
    parser.add_argument("--css", type=str, help="Add css as [url|path|content]")
    parser.add_argument("--config", type=str, help="Path to the config file.")
    parser.add_argument("--emulateMedia", type=str, choices=['screen', 'print'], default='screen', help="Type of emulate media.")
    parser.add_argument("--qr", type=int, help="Add QR after link.")

    parser.add_argument("--merge", nargs='+', help="Merge PDF files into an output file")
    parser.add_argument("--extract", help="Extract pages from a PDF file (e.g., '1-3')")
    parser.add_argument("--split", action='store_true', help="Split a PDF file at specified pages")
    parser.add_argument("--extract-text", action='store_true', help="Extract text from a PDF file") # TODO
    parser.add_argument("--redact-text", nargs='+', help="Redact text in the PDF files")
    parser.add_argument("--meta", type=Parser.meta, help="Add meta in the format '/Key: Value'. Example: '/Title Document'")
    parser.add_argument("--remove-meta", action='store_true', help="Remove meta from a PDF file")
    parser.add_argument("--select-meta", nargs='+', help="Selectively remove meta fields")

    parser.add_argument("--verbose", action="store_true", help="Print parameters when verbose.")

    parser.add_argument("--encrypto", action="store_true", help="Encrypt PDF file.")
    parser.add_argument("--autogen", action="store_true", help="Automatically generate password for encryption.")
    parser.add_argument("--decrypto", action="store_true", help="Decrypt PDF file.")
    parser.add_argument("--regen", action="store_true", help="Regenerate and send recovery key for decryption.")

    parser.add_argument("--webui", action="store_true", help="[WIP] Launch Gradio web UI.")

    parser.add_argument("--share", choices=["mega", "google_drive", "slideshare"], help="Specify the cloud service to use for sharing")

    parser.usage = '''
        [Examples]

        Generate:
            From local file:
                xdoc -i in_*.html --verbose

            From remote file:
                set gist https://gist.githubusercontent.com/.../gistfile.md
                xdoc -i $gist

            Styling:
                xdoc -i in_*.html --scale 0.8 --css ' h1 { color: red }' --margin 'top bottom 50mm left right 20mm'

            Load Config:
                xdoc --config config.yaml

        Modify:
            Extract pages:
                xdoc -i in.pdf --extract 2-3

            meta
                xdoc -i in.md --meta '/Title Untitled'

        Encrypt/Decrypto:
            xdoc -i demo.pdf -o demo.encrypted.pdf --encrypto --autogen
            xdoc -i demo.encrypted.pdf -o demo.decrypted.pdf --decrypto

        Share:
            python -i $gist -o out.pdf --share mega


        [Units]

        You can use the following units for width, height, and margin options:
        - px (pixels) - Default unit., in (inches)., cm (centimeters)., mm (millimeters).


        [Formats]

        The following paper formats are available for the format option:
        - Letter, Legal, Tabloid, Ledger, A0, A1, A2, A3, A4 (default), A5


        [PDF meta]

        meta format: "/Key Value" pairs separated by space.
        - /Title, /Author, /Subject, /Author, /Keyword, /Creator, /Producer, /CreateDate,  /ModDate, /Trapped
        '''

    args = parser.parse_args()
    params = {}

    if args.config:
        with open(args.config, "r") as config_file:
            config_data = yaml.safe_load(config_file)
            params = config_data.get("params", {})

    for key, value in vars(args).items():
        if value is not None and value is not False:
            params[key] = value

    if args.verbose:
        console.print('---')
        for key, value in params.items():
            console.print(f'{key}: {value}---\n\n')

    if args.input:
        ts, tmp_dir = Interface.get_tempdir()

        tmp_dir = f'{tmp_dir}/{ts}'
        tmp_file = f'{tmp_dir}/{ts}'
        os.makedirs(tmp_dir, exist_ok=True)

        input_files = Parser.input_txt(args.input)

        params.update({'input': input_files})
        params.update({'tmp_dir': tmp_dir, 'tmp_file': f'{tmp_dir}/{ts}'})

        processing_files = []

        if args.output:
            params.update({'output': args.output})
            result_file = args.output[0]

        else:
            cwd = os.getcwd()
            base = os.path.splitext(os.path.basename(input_files[0]))[0] + f'.pdf'
            abspath_dir = os.path.abspath(base)
            result_file = os.path.join(abspath_dir, base)
            params.update({'output': result_file})
            console.print(f'[i] Output will: {result_file}')

        for index, file in enumerate(input_files):
            mime_type, _ = mimetypes.guess_type(file)
            if mime_type and mime_type.startswith('application'):
                processing_files.append(result_file)
            elif mime_type and mime_type.startswith('text'):
                processing_files.append(result_file)
            else:
                console.print(f"[!] Unsupported file type for {file}")
        console.print()

        if args.qr:
            qr_params = {
                'version': 1,
                'error_correction': qrcode.constants.ERROR_CORRECT_L,
                'box_size': params['qr'],
                'border': 1,
            }
            await ImageProcessor.add_qr(input_files, qr_params, tmp_dir, tmp_file)

        pdf_proc = PdfProcessor(params)

        processing_file = await pdf_proc.gen()

        if args.merge:
            choosen_file = choose_file(processing_files)
            pdf_proc.merge(args.merge[1:], args.merge[0])

        if args.extract:
            choosen_file = choose_file(processing_files)
            page_range = args.extract
            start_page, end_page = map(int, page_range.split('-'))
            pdf_proc.write(choosen_file, result_file, start_page, end_page)

        if args.split:
            choosen_file = choose_file(processing_files)
            pdf_proc.split(choosen_file, result_file, 1)

        if args.extract_text:
            choosen_file = choose_file(processing_files)
            pdf_proc.extract_text(choosen_file, result_file)

        if args.remove_meta:
            choosen_file = choose_file(processing_files)
            pdf_proc.remove_meta(choosen_file, result_file)

        if args.select_meta:
            choosen_file = choose_file(processing_files)
            pdf_proc.select_meta(choosen_file, result_file, args.select_meta)

        if args.encrypto and args.decrypto:
            console.print("[!] Cannot specify both --encrypto and --decrypto simultaneously.")

        elif args.encrypto:
            in_ = processing_file or choose_file(args.input)
            pwd = Encryptor.gen_pass() if args.autogen else getpass.getpass("Enter pwd: ")
            pwd_confirm = None
            while pwd != pwd_confirm and not args.autogen:
                pwd_confirm = getpass.getpass("Confirm pwd: ")
                if pwd != pwd_confirm:
                    console.print("[!] pwds do not match.")
            pdf_proc.encrypt(in_, pwd)
            console.print(f"[*] Encrypted file Output to: {args.output}")

        elif args.decrypto:
            if not args.regen:
                pwd = getpass.getpass("Enter pwd: ")
                pdf_proc.decrypt(pwd)
            else:
                console.print("--regen: This feature is WIP.")

        if args.share:
            if f"{args.share.upper()}_ID" in os.environ and f"{args.share.upper()}_PASS" in os.environ:
                username = os.getenv[f"{args.share.upper()}_ID"]
                password = os.getenv[f"{args.share.upper()}_PASS"]
            else:
                console.print(f"Enter {args.share} ID:")
                username = input()
                password = getpass.getpass(f"Enter {args.share} password:")

            WebAPI.share_file(args.share, username, password, result_file)


    elif args.webui:
        asyncio.create_task(webui(params, parser.usage))

    else:
        logger.setLevel(logging.NOTSET)
        parser.print_help()


if __name__ == "__main__":
    asyncio.run(main())
