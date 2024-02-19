import re
import os
import docx
import openai
import shutil
import requests

from docx.shared import Inches
from docx2pdf import convert
from bruhcolor import bruhcolored as bc
from dotenv import load_dotenv

from bruhbook.bruhbookerrors import (
    ApiKeyNotFoundError,
    InvalidParameterValueError,
    UnhandledExceptionError,
    MissingParameterError,
)


class BruhBook:
    def __init__(
        self,
        api_key: str | None = None,
        create_cover_image: bool = False,
        wipe_files: bool = False,
    ):
        """Constructor for BruhBook

        Args:
            api_key (str | None, optional): OpenAI API key. Defaults to None.
            create_cover_image (bool, optional): To make api calls to OpenAI to
            generate a cover image for the book. Defaults to False.

        Raises:
            ApiKeyNotFoundError: Error for not finding an OpenAI API key
        """

        if api_key:
            self.__api_key = api_key
        elif ak := os.getenv("OPENAI_API_KEY"):
            self.__api_key = ak
        else:
            raise ApiKeyNotFoundError(None)

        self.__openai_client = openai.OpenAI(api_key=self.__api_key)
        self.__generation_model = "gpt-3.5-turbo-0125"
        self.__image_model = "dall-e-3" if self.__check_for_dalle_3() else "dall-e-2"
        self.__create_cover_image = create_cover_image
        self.__show_image_prompt = True
        self.__story_outline = None
        self.__wipe_files = wipe_files
        self.__prompts = {
            "image_prompt_creator_prompt": {
                "system": "\nYou are an expert image prompt creator. You create prompts for people that will be used to create images with tools like DALL-E-3 or MidJourney.\nThese prompts are used to create cover art for books. You will be provided what the book is about and the target audience.\n- You will be given what the story is about surrounded in four hashtags (####) (e.g. #### [Book info here] ####)\n- You will be given the target audience of the book surrounded in four percent symbols (%%%%) (e.g. %%%% 5 - 7 year olds %%%%)\n\nReturn back a prompt that can be used with DALL-E-2 or Midjourney that would create a cover art for this book. Make sure it is detailed! Make sure to focus on the outcome of the entire picture, not individual attributes of the image.\n",
                "user": "\nBook Info: #### ${book_info} ####\nTarget Audience: %%%% ${target_audience} %%%%\n[Prompt for iamge creation tool here]\n",
                "variables": ["book_info", "target_audience"],
            },
            "title_creator_prompt": {
                "system": "You are an expert book title maker. You are given some info about the book, create a simple 5 word or less title.",
                "user": "%{story_type}\n[Title Here]",
                "variables": ["story_type"],
            },
            "partial_story_generator": {
                "system": "\nYou are an expert partial story generator. Your task is to creatively expand upon provided story details, crafting a coherent and engaging single, multi-paragraph page to add to the story.\n\nPlease follow these guidelines:\n\n- Story Title/Purpose: Identified by four hashtags (e.g., #### Story Title ####).\n- Target Audience: Specified within four @ symbols (e.g., @@@@ Adults @@@@).\n- Chapter Title: Highlighted by four dollar signs (e.g., $$$$ Chapter Title $$$$).\n- Page Overview: Enclosed in four asterisks (e.g., **** Page Details ****).\n- Optional - Previous Page Knowledge: Provided within four exclamation points (e.g., !!!! Previous Page Summary !!!!).\n- Last Chapter Indicator: A boolean flag (True/False) to indicate if it's the last chapter.\n- Last Page Indicator: A boolean flag (True/False) to indicate if it's the last page of the story.\n\nYour Task:\n\n- Using the provided details, create the next page of the story.\n- Ensure continuity with previous content if available.\n- Aim for a length of approximately [specify word/paragraph count].\n- Be creative while adhering to the guidelines.\n- If 'Last Chapter' is True, ensure the chapter concludes appropriately.\n- If 'Last Page' is True, bring the story to a satisfying conclusion.\n\nFor example:\n#### The Lost City ####\n@@@@ Young Adults @@@@\n$$$$ The Secret Door $$$$\n**** Write about the protagonist discovering a hidden door in the ruins ****\n!!!! \nIn the previous page, the protagonist was exploring ancient ruins. \n!!!!\nLast Chapter: True\nLast Page of Chapter: False\n[Your story content here]\n\nNote: Do not include the chapter title in your story content.\n",
                "user": "\n#### ${story_title} ####\n@@@@ ${target_audience} @@@@\n$$$$ ${chapter_title} $$$$\n**** ${page_to_create} ****\n!!!!\n${last_paragraph}\n!!!!\nLast Chapter: ${last_chapter}\nLast Page of Chapter: ${last_page}\n[Your story content here]\n",
                "variables": [
                    "story_title",
                    "target_audience",
                    "chapter_title",
                    "page_to_create",
                    "last_paragraph",
                    "last_chapter",
                    "last_page",
                ],
            },
            "story_outline_generator": {
                "system": "\nYou are an expert story outliner.\nYou will be provided a type of story and target audience.\nYour task is to take these two items and create a story outline that has 5 main plot points.\n\nYou will be given the story type surrounded in four hashtags (####). As an example, #### <story type> ####.    \nYou will be given the target audience surrounded in four dollar signs ($$$$). As an example, $$$$ <story type> $$$$.\n\nRules:\n- You MUST generate a rich story outline with 3 to 5 main plot points.\n- You MUST lead main points with a single dash.\n- You MUST provide 5 to 10 sub points for each main point that explains the story line for the main point.\n- You MUST lead sub points with two dashes.\n\n\nExample Input:\nStory Type:\n####\n<story type here>\n####\n\nTarget Audience:\n$$$$\n<target audience here>\n$$$$\n\n\nExample Output:\n- Point one\n-- sub point 1-1\n-- sub point 1-2\n-- sub point 1-3\n. . . \n",
                "user": "\nStory Type:\n####\n${story_type}\n####\n\nTarget Audience:\n$$$$\n${target_audience}\n$$$$\n",
                "variables": ["story_type", "target_audience"],
            },
        }

        self.__home_directory = os.getcwd()

        self.__check_for_folders()

    def __check_for_folders(self) -> None:
        """_summary_"""

        stories_path = self.__home_directory + "/stories"
        if not os.path.exists(stories_path):
            os.mkdir(stories_path)

    def set_generation_model(self, model: str):
        self.__generation_model = model

    def set_show_image_prompt(self, show_image_prompt: bool) -> None:
        self.__show_image_prompt = show_image_prompt

    def __check_for_dalle_3(self) -> bool:
        return "dall-e-3" in [model.id for model in self.__openai_client.models.list()]

    def __openai_prompter(
        self,
        system_prompt: str,
        user_prompt: str,
        api_type: str = "generate",
        image_path: str | None = None,
        model: str | None = None,
    ):
        """_summary_

        Args:
            system_prompt (str): _description_
            user_prompt (str): _description_
            api_type (str, optional): _description_. Defaults to "generate".
            image_path (str | None, optional): _description_. Defaults to None.
            model (str | None, optional): _description_. Defaults to None.

        Raises:
            UnhandledExceptionError: _description_
            MissingParameterError: _description_
            InvalidParameterValueError: _description_

        Returns:
            _type_: _description_
        """

        if api_type == "generate":
            try:
                return (
                    self.__openai_client.chat.completions.create(
                        messages=[
                            {"role": "system", "content": system_prompt},
                            {"role": "user", "content": user_prompt},
                        ],
                        temperature=0.5,
                        max_tokens=1500,
                        model=model if model else self.__generation_model,
                    )
                    .choices[0]
                    .message.content
                )
            except Exception as e:
                raise UnhandledExceptionError(str(e))

        elif api_type == "image":
            if image_path:
                try:
                    response = self.__openai_client.images.generate(
                        model=model if model else self.__image_model,
                        prompt=user_prompt,
                        size="1024x1024",
                        quality="standard",
                        n=1,
                    )
                    image_url = response.data[0].url
                    image_data = requests.get(image_url).content
                    with open(image_path + "/cover_art.png", "wb") as image_file:
                        image_file.write(image_data)
                except Exception as e:
                    raise UnhandledExceptionError(str(e))
            else:
                raise MissingParameterError("image_path")
        else:
            raise InvalidParameterValueError(
                value=api_type, parameter="api_type", options=["generate", "image"]
            )

    def improve_image_prompt(self, book_info: str, target_audience: str) -> str:
        try:
            improved_prompt = self.chat(
                system_prompt=self.__prompts["image_prompt_creator_prompt"]["system"],
                user_prompt=self.__prompts["image_prompt_creator_prompt"]["user"]
                .replace(
                    self.__prompts["image_prompt_creator_prompt"]["variables"][0],
                    book_info,
                )
                .replace(
                    self.__prompts["image_prompt_creator_prompt"]["variables"][1],
                    target_audience,
                ),
                model="gpt-4",
            )
        except Exception as e:
            raise UnhandledExceptionError(str(e))

        return improved_prompt

    def create_cover_art(
        self,
        story_type: str,
        target_audience: str,
        image_path: str,
        model: str | None = None,
        show_prompt: bool = False,
    ) -> None:
        cover_art_prompt = self.improve_image_prompt(
            book_info=story_type, target_audience=target_audience
        )

        if self.__show_image_prompt:
            print(bc("Prompt for cover art: " + cover_art_prompt, color=202))

        self.__image(user_prompt=cover_art_prompt, image_path=image_path, model=model)

    def chat(
        self, system_prompt: str, user_prompt: str, model: str | None = None
    ) -> str:
        """_summary_

        Args:
            system_prompt (str): _description_
            user_prompt (str): _description_

        Returns:
            str: _description_
        """

        return self.__openai_prompter(
            system_prompt=system_prompt,
            user_prompt=user_prompt,
            api_type="generate",
            model=model,
        )

    def __image(
        self, user_prompt: str, image_path: str, model: str | None = None
    ) -> None:
        """_summary_

        Args:
            user_prompt (str): _description_
            image_path (str): _description_
            model (str | None, optional): _description_. Defaults to None.
        """

        self.__openai_prompter(
            system_prompt=None,
            user_prompt=user_prompt,
            api_type="image",
            image_path=image_path,
            model=model,
        )

    def get_client(self) -> openai.OpenAI:
        return self.__openai_client

    def generate_story_outline(
        self, story_type: str, target_audience: str
    ) -> dict[str, list[str]]:
        openai_story_outline = self.chat(
            system_prompt=self.__prompts["story_outline_generator"]["system"],
            user_prompt=self.__prompts["story_outline_generator"]["user"]
            .replace(
                self.__prompts["story_outline_generator"]["variables"][0], story_type
            )
            .replace(
                self.__prompts["story_outline_generator"]["variables"][1],
                target_audience,
            ),
        ).split("\n")

        outline = [val for val in openai_story_outline if val.strip() != ""]
        story_outline = {}
        current_chapter = ""

        for val in outline:
            if re.match(r"\-\s.*", val):
                current_chapter = val.split("- ")[1]
                story_outline[current_chapter] = []
            elif re.match(r"\-\-\s.*", val):
                story_outline[current_chapter].append(val.split("-- ")[1])
            else:
                print("ERROR")

        self.__story_outline = story_outline

    def save_to_pdf(self, base_path: str, story_type: str) -> None:

        main_title = self.chat(
            system_prompt=self.__prompts["title_creator_prompt"]["system"],
            user_prompt=self.__prompts["title_creator_prompt"]["user"].replace(
                self.__prompts["title_creator_prompt"]["variables"][0], story_type
            ),
        )

        main_title_sub = re.sub(r"[^A-Za-z0-9 ]", "", main_title).replace(" ", "_")

        print(bc(f"Creating '{base_path}/{main_title_sub}.pdf'", color=72))
        doc = docx.Document()

        title = os.path.basename(base_path)
        doc.add_paragraph(main_title, "Title")
        if self.create_cover_art:
            doc.add_picture(f"{base_path}/cover_art.png", width=Inches(6))
        doc.add_page_break()
        for chapter_folder in os.listdir(base_path):
            chapter_path = os.path.join(base_path, chapter_folder)
            if os.path.isdir(chapter_path):
                doc.add_heading(
                    re.sub(r"[^A-Za-z0-9 ]", "", chapter_folder.replace("_", " ")),
                    level=1,
                )
                for page_file in os.listdir(chapter_path):
                    page_path = os.path.join(chapter_path, page_file)
                    if os.path.isfile(page_path) and page_file.endswith(".txt"):
                        with open(page_path, "r") as file:
                            doc.add_paragraph("".join(file.readlines()[2:]))
            doc.add_page_break()

        doc.save(f"{base_path}/{main_title_sub}.docx")
        try:
            convert(
                f"{base_path}/{main_title_sub}.docx",
                f"{base_path}/{main_title_sub}.pdf",
            )
            os.remove(f"{base_path}/{main_title_sub}.docx")
        except Exception as e:
            print(e)

        if self.__wipe_files:
            print(
                bc(
                    f"Removing '{base_path}/<chapter folders>', keeping PDF and Cover Art",
                    color=196,
                )
            )
            for path_name in os.listdir(f"{base_path}"):
                if os.path.isdir(f"{base_path}\{path_name}"):
                    shutil.rmtree(f"{base_path}\{path_name}")

    def __chapter_page_generator(
        self,
        story_type: str,
        target_audience: str,
        chapter_title: str,
        chapter_point: str,
        previous_knowledge: str,
        completed_chapters: str,
        last_paragraph: str,
        chapter_save_path: str,
        page_number: int,
        last_chapter: bool = False,
        last_page: bool = False,
    ) -> str:
        if previous_knowledge == "":
            previous_knowledge = "No Previous Knowledge."
        if completed_chapters == "":
            completed_chapters = "No Previously Completed Chapters."
        if last_paragraph == "":
            last_paragraph = "No Last Paragraph - Start of a new chatper."

        filled_user_prompt = (
            self.__prompts["partial_story_generator"]["user"]
            .replace("${story_title}", story_type)
            .replace("${target_audience}", target_audience)
            .replace("${chapter_title}", chapter_title)
            .replace("${page_to_create}", chapter_point)
            .replace("${previous_knowledge}", previous_knowledge)
            .replace("${completed_chapters}", completed_chapters)
            .replace("${last_paragraph}", last_paragraph)
            .replace("${last_chapter}", "is" if last_chapter else "is not")
            .replace("${last_page}", "is" if last_page else "is not")
        )

        new_page = self.chat(
            system_prompt=self.__prompts["partial_story_generator"]["system"],
            user_prompt=filled_user_prompt,
        )

        last_paragraph = [val for val in new_page.split("\n") if val.strip() != ""][-1]

        page_path = f"{chapter_save_path}/page_{page_number}.txt"

        with open(page_path, "w") as file:
            print(bc(f"Creating Chapter Page: '{page_path}'", color=87))
            file.write(
                f"Story Type: {story_type}\n"
                + f"Chapter Title: {chapter_title}\n\n"
                + new_page
                + "\n"
            )

        return (
            previous_knowledge
            + f"Completed Page: {chapter_title} - Page {page_number} - {chapter_point}\n",
            last_paragraph,
        )

    def story_generator(self, story_type: str, target_audience: str) -> None:
        cleaned_story_type = re.sub(r"[^A-Za-z0-9 ]", "", story_type).replace(" ", "_")

        if not os.path.exists("./stories"):
            print(bc("Creating 'stories' base folder", color=82))
            os.mkdir("./stories")

        if not os.path.exists(f"./stories/{cleaned_story_type}"):
            print(bc(f"Creating 'stories/{cleaned_story_type}' story folder", color=72))
            os.mkdir(f"./stories/{cleaned_story_type}")

        base_path = f"./stories/{cleaned_story_type}"

        if self.__create_cover_image:
            self.create_cover_art(
                story_type=story_type,
                target_audience=target_audience,
                image_path=base_path,
            )

        chapter_number = 0
        completed_chapters = ""

        for chapter_num, chapter_val in enumerate(self.__story_outline.items()):
            last_chapter = chapter_num == len(self.__story_outline) - 1
            chapter, chapter_points = chapter_val
            chapter_number += 1
            previous_chapter_knowledge = ""
            last_paragraph = ""
            cleaned_chapter_name = f"{chapter_number}_" + re.sub(
                r"[^A-Za-z0-9 ]", "", chapter
            ).replace(" ", "_")
            chapter_page_save_path = (
                f"./stories/{cleaned_story_type}/{cleaned_chapter_name}"
            )
            print(bc(f"Creating Chapter Folder '{chapter_page_save_path}'", color=206))
            os.mkdir(chapter_page_save_path)
            for idx, point in enumerate(chapter_points):
                last_page = idx == len(chapter_points) - 1
                previous_chapter_knowledge, last_paragraph = (
                    self.__chapter_page_generator(
                        story_type=story_type,
                        target_audience=target_audience,
                        chapter_title=chapter,
                        chapter_point=point,
                        previous_knowledge=previous_chapter_knowledge,
                        completed_chapters=completed_chapters,
                        last_paragraph=last_paragraph,
                        chapter_save_path=chapter_page_save_path,
                        page_number=idx,
                        last_chapter=last_chapter,
                        last_page=last_page,
                    )
                )
            completed_chapters += f"Previously Completed Chapter Title - {chapter}\n"

        self.save_to_pdf(base_path=base_path, story_type=cleaned_story_type)

    def generate_story(self, story_type: str, target_audience: str | None = "Anyone"):
        self.generate_story_outline(
            story_type=story_type, target_audience=target_audience
        )

        self.story_generator(
            story_type=story_type,
            target_audience=target_audience,
        )
